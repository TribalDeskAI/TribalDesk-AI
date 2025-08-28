import streamlit as st
from io import BytesIO
from docx import Document
from openai import OpenAI

# ----------------------
# Helpers
# ----------------------

def get_client():
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def assemble_text(org, project, summary, need, goals, methods, evaln, budget, sustain, sovereignty):
    return f"""
# {project or 'Untitled Project'} — Proposal Draft
**Organization:** {org or 'N/A'}

## Summary
{summary}

## Needs Statement
{need}

## Goals & Objectives
{goals}

## Methods / Activities
{methods}

## Evaluation
{evaln}

## Budget
{budget}

## Sustainability
{sustain}

## Alignment with Tribal Sovereignty & Culture
{sovereignty}
"""

# ----------------------
# Streamlit UI
# ----------------------

st.title("TribalDeskAI Proposal Generator")

# Form inputs
with st.form("proposal_form"):
    org = st.text_input("Organization")
    project = st.text_input("Project Title")
    summary = st.text_area("Summary")
    need = st.text_area("Needs Statement")
    goals = st.text_area("Goals & Objectives")
    methods = st.text_area("Methods / Activities")
    evaln = st.text_area("Evaluation")
    budget = st.text_area("Budget")
    sustain = st.text_area("Sustainability")
    sovereignty = st.text_area("Alignment with Tribal Sovereignty & Culture")
    
    use_ai = st.checkbox("Enhance with AI", value=True)
    submitted = st.form_submit_button("Generate Proposal")

if submitted:
    base_text = assemble_text(org, project, summary, need, goals, methods, evaln, budget, sustain, sovereignty)

    if use_ai:
        client = get_client()
        try:
            prompt = (
                "Improve and expand the following grant proposal sections for clarity, "
                "structure, and persuasive impact. Keep culturally respectful language "
                "and align with Tribal sovereignty. Return clean Markdown.\n\n" + base_text
            )
            resp = client.chat.completions.create(
                model=st.secrets.get("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role":"user","content": prompt}]
            )
            final_md = resp.choices[0].message.content
        except Exception as e:
            st.error(f"OpenAI error: {e}")
            final_md = base_text
    else:
        final_md = base_text

    # ----------------------
    # Preview
    # ----------------------
    st.markdown("---")
    st.subheader("Draft Preview")
    st.markdown(final_md)

    # ----------------------
    # DOCX Export
    # ----------------------
    doc = Document()
    for line in final_md.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        else:
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    st.download_button(
        "Download .docx",
        data=bio.getvalue(),
        file_name=f"{project or 'proposal'}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Sidebar nav
st.sidebar.page_link("streamlit_app.py", label="🏠 Home")
