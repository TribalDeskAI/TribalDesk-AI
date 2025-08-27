# TribalDeskAI — Fresh Rebuild Starter Kit

This pack gives you a clean, working, monetizable Streamlit app with:

* A professional **Landing Page**
* **Chat Assistant** powered by OpenAI
* **Proposal Generator** (AI-assisted) with document downloads
* **Grant Finder** tracker (CSV storage; optional Google Sheets later)
* **Email capture** saved to `data/emails.csv`
* **Payhip** store buttons/embeds
* Safe **secrets** handling and a one-click deploy flow

Use the exact file/folder names below.

---

## 📁 Repository Structure

```
tribaldesk-ai/
├─ streamlit_app.py                 # Landing page (home)
├─ pages/
│  ├─ 1_Chat_Assistant.py
│  ├─ 2_Proposal_Generator.py
│  └─ 3_Grant_Finder.py
├─ utils/
│  ├─ openai_client.py
│  └─ storage.py
├─ data/
│  └─ company_context.md            # Business/landing context injected into AI
├─ requirements.txt
├─ .gitignore
├─ README.md
└─ .streamlit/
   └─ secrets.toml                  # (You create this on Streamlit Cloud, not in repo)
```

---

## 🔹 `streamlit_app.py` (Landing Page)

```python
import streamlit as st
from utils.storage import ensure_data_dir, save_email

# ===== Config =====
COMPANY = "TribalDesk AI"
TAGLINE = "Empowering Tribal sovereignty with AI-driven grants, policies, and training."
CONTACT_EMAIL = st.secrets.get("CONTACT_EMAIL", "info@tribaldeskai.com")
PAYHIP_STORE_URL = st.secrets.get("PAYHIP_STORE_URL", "")

st.set_page_config(page_title=COMPANY, page_icon="🌐", layout="wide")
ensure_data_dir()

# ===== Header / Hero =====
col_logo, col_cta = st.columns([3, 1])
with col_logo:
    st.markdown(f"### {COMPANY}")
    st.markdown(f"**{TAGLINE}**")
with col_cta:
    st.page_link("pages/1_Chat_Assistant.py", label="🚀 Open Chat Assistant", icon="💬")
    st.page_link("pages/2_Proposal_Generator.py", label="📝 Proposal Generator", icon="📝")
    st.page_link("pages/3_Grant_Finder.py", label="🔍 Grant Finder", icon="🔎")

st.markdown("---")

# ===== Value / Solutions =====
st.header("Our Solutions")
col1, col2, col3 = st.columns(3)
with col1:
    st.subheader("AI Writing Tools")
    st.write("Grant proposals, policies, reports — faster and clearer.")
    st.page_link("pages/2_Proposal_Generator.py", label="Try Proposal Generator")
with col2:
    st.subheader("Funding Platform")
    st.write("Track and manage funding opportunities tailored for Tribal communities.")
    st.page_link("pages/3_Grant_Finder.py", label="Open Grant Finder")
with col3:
    st.subheader("Consulting & Training")
    st.write("Peacekeeping mediation programs, strategic planning, workshops.")
    if PAYHIP_STORE_URL:
        st.link_button("Browse Templates & Trainings (Payhip)", PAYHIP_STORE_URL)

# ===== Why Choose Us =====
st.markdown("---")
st.header("Why Choose TribalDesk AI")
st.write("- Unmatched Tribal focus and cultural competency.\n- Streamlined efficiency with AI.\n- Accessible, affordable, and practical tools.")

# ===== Who We Serve =====
st.markdown("---")
st.header("Who We Serve")
st.write("- Tribal Governments & Courts\n- Tribal Nonprofits\n- Native-owned Small Businesses\n- Grant Writers & Consultants")

# ===== Testimonials (placeholders) =====
st.markdown("---")
st.header("What Our Partners Say")
st.info("\"TribalDesk AI helped us cut our grant writing time by 70%.\" — Chief Admin, Example Nation")
st.info("\"Found funding we never knew existed.\" — CEO, Native-owned Business")

# ===== Email Capture =====
st.markdown("---")
st.header("Stay in the loop")
with st.form("email_capture"):
    email = st.text_input("Enter your email")
    interest = st.text_input("What are you most interested in? (optional)")
    submitted = st.form_submit_button("Subscribe")
    if submitted:
        if email and "@" in email:
            save_email(email, interest)
            st.success(f"Thanks! We'll reach out at {email}.")
        else:
            st.error("Please enter a valid email.")

# ===== Footer / Monetization =====
st.markdown("---")
col_f1, col_f2 = st.columns(2)
with col_f1:
    st.caption(f"© 2025 {COMPANY} · {CONTACT_EMAIL}")
with col_f2:
    if PAYHIP_STORE_URL:
        st.link_button("🛒 Visit our Payhip Store", PAYHIP_STORE_URL)
```

---

## 🔹 `pages/1_Chat_Assistant.py`

```python
import streamlit as st
from utils.openai_client import get_client

st.set_page_config(page_title="Chat Assistant", page_icon="💬", layout="wide")

st.title("💬 TribalDesk AI — Chat Assistant")
st.caption("Ask about grants, proposals, mediation programs, startup steps, and more.")

# Load context to steer the model
from pathlib import Path
ctx_path = Path("data/company_context.md")
company_context = ctx_path.read_text(encoding="utf-8") if ctx_path.exists() else ""

# Model selection (defaults)
model_default = st.secrets.get("OPENAI_MODEL", "gpt-4o-mini")
model = st.sidebar.selectbox("Model", [model_default, "gpt-4o", "gpt-4o-mini", "gpt-4.1-mini", "gpt-3.5-turbo"], index=0)

# Initialize
if "chat_history" not in st.session_state:
    st.session_state.chat_history = [
        {"role": "system", "content": (
            "You are TribalDesk AI, an expert assistant for Tribal governments, nonprofits, and Native-owned businesses. "
            "Be clear, concise, and culturally respectful. Use the provided company context when helpful.\n\n" + company_context
        )}
    ]

client = get_client()

# Chat UI
for msg in [m for m in st.session_state.chat_history if m["role"] != "system"]:
    if msg["role"] == "user":
        st.markdown(f"**You:** {msg['content']}")
    else:
        st.markdown(f"**Assistant:** {msg['content']}")

user_input = st.text_area("Type your question…", placeholder="How do I structure a DOJ CTAS proposal?", height=100)
col_send, col_clear = st.columns([1,1])
if col_send.button("Send"):
    if user_input.strip():
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        try:
            stream = client.chat.completions.create(
                model=model,
                messages=st.session_state.chat_history,
                stream=True
            )
            full = ""
            placeholder = st.empty()
            for chunk in stream:
                delta = chunk.choices[0].delta.get("content", "")
                full += delta
                placeholder.markdown(f"**Assistant:** {full}")
            st.session_state.chat_history.append({"role": "assistant", "content": full})
        except Exception as e:
            st.error(f"OpenAI error: {e}")
if col_clear.button("Clear Chat"):
    st.session_state.chat_history = st.session_state.chat_history[:1]  # keep system
    st.experimental_rerun()

st.sidebar.page_link("streamlit_app.py", label="🏠 Home")
```

---

## 🔹 `pages/2_Proposal_Generator.py`

```python
import streamlit as st
from utils.openai_client import get_client
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Proposal Generator", page_icon="📝", layout="wide")
st.title("📝 Grant Proposal Generator")
st.caption("Draft a strong proposal fast. Fill fields and optionally let AI expand your text.")

with st.form("proposal_form"):
    org = st.text_input("Organization name")
    project = st.text_input("Project title")
    summary = st.text_area("1) Project Summary / Abstract")
    need = st.text_area("2) Needs Statement (Tribal context, data, community voice)")
    goals = st.text_area("3) Goals & Objectives")
    methods = st.text_area("4) Methods / Activities (timeline, partners)")
    evaln = st.text_area("5) Evaluation (metrics, feedback, learning)")
    budget = st.text_area("6) Budget (categories, match, justification)")
    sustain = st.text_area("7) Sustainability (post-award, capacity)")
    sovereignty = st.text_area("8) Alignment with Tribal Sovereignty & Culture")

    use_ai = st.checkbox("Use AI to expand sections with improvements")
    submitted = st.form_submit_button("Generate Draft")

if submitted:
    def assemble_text(**k):
        return f"""
# {k['project']} — Proposal Draft
Organization: {k['org']}

## Summary
{k['summary']}

## Needs Statement
{k['need']}

## Goals & Objectives
{k['goals']}

## Methods / Activities
{k['methods']}

## Evaluation
{k['evaln']}

## Budget
{k['budget']}

## Sustainability
{k['sustain']}

## Alignment with Tribal Sovereignty & Culture
{k['sovereignty']}
"""

    if use_ai:
        client = get_client()
        try:
            base_text = assemble_text(org=org, project=project, summary=summary, need=need, goals=goals, methods=methods, evaln=evaln, budget=budget, sustain=sustain, sovereignty=sovereignty)
            prompt = (
                "Improve and expand the following grant proposal sections for clarity, structure, and persuasive impact. "
                "Keep culturally respectful language and align with Tribal sovereignty. Return clean Markdown.\n\n" + base_text
            )
            resp = client.chat.completions.create(
                model=st.secrets.get("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role":"user","content": prompt}]
            )
            final_md = resp.choices[0].message.content
        except Exception as e:
            st.error(f"OpenAI error: {e}")
            final_md = assemble_text(org=org, project=project, summary=summary, need=need, goals=goals, methods=methods, evaln=evaln, budget=budget, sustain=sustain, sovereignty=sovereignty)
    else:
        final_md = assemble_text(org=org, project=project, summary=summary, need=need, goals=goals, methods=methods, evaln=evaln, budget=budget, sustain=sustain, sovereignty=sovereignty)

    st.markdown("---")
    st.subheader("Draft Preview")
    st.markdown(final_md)

    # Download as .docx
    doc = Document()
    for line in final_md.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    st.download_button("Download .docx", data=bio.getvalue(), file_name=f"{project or 'proposal'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.sidebar.page_link("streamlit_app.py", label="🏠 Home")
```

---

## 🔹 `pages/3_Grant_Finder.py`

```python
import streamlit as st
import pandas as pd
from utils.storage import ensure_data_dir, load_csv, save_csv
from pathlib import Path

st.set_page_config(page_title="Grant Finder", page_icon="🔎", layout="wide")
st.title("🔎 Grant Finder & Tracker")
st.caption("Manually add and track grants. Advanced scraping/search can be added later.")

ensure_data_dir()
DATA_FILE = Path("data/grants.csv")

# Load or init
if DATA_FILE.exists():
    df = load_csv(DATA_FILE)
else:
    df = pd.DataFrame(columns=["Title","Funder","Link","Deadline","Amount","Notes","Status"])
    save_csv(DATA_FILE, df)

st.subheader("Add a grant")
with st.form("add_grant"):
    title = st.text_input("Title")
    funder = st.text_input("Funder")
    link = st.text_input("Link (URL)")
    deadline = st.date_input("Deadline")
    amount = st.text_input("Amount (e.g., $100,000)")
    notes = st.text_area("Notes")
    status = st.selectbox("Status", ["To Review","Drafting","Submitted","Awarded","Declined"])
    add = st.form_submit_button("Add")

if add and title:
    new = {"Title": title, "Funder": funder, "Link": link, "Deadline": deadline, "Amount": amount, "Notes": notes, "Status": status}
    df = pd.concat([df, pd.DataFrame([new])], ignore_index=True)
    save_csv(DATA_FILE, df)
    st.success("Grant added.")

st.markdown("---")
st.subheader("Your Grants")
st.dataframe(df, use_container_width=True)

csv_bytes = df.to_csv(index=False).encode("utf-8")
st.download_button("Export CSV", data=csv_bytes, file_name="grants.csv", mime="text/csv")

st.sidebar.page_link("streamlit_app.py", label="🏠 Home")
```

---

## 🔹 `utils/openai_client.py`

```python
import streamlit as st
from openai import OpenAI

def get_client() -> OpenAI:
    api_key = st.secrets.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("Missing OPENAI_API_KEY in Streamlit Secrets.")
    return OpenAI(api_key=api_key)
```

---

## 🔹 `utils/storage.py`

```python
import streamlit as st
import pandas as pd
from pathlib import Path

DATA_DIR = Path("data")

def ensure_data_dir():
    DATA_DIR.mkdir(parents=True, exist_ok=True)

def save_email(email: str, interest: str = ""):
    ensure_data_dir()
    f = DATA_DIR / "emails.csv"
    if f.exists():
        df = pd.read_csv(f)
    else:
        df = pd.DataFrame(columns=["email","interest"])
    if email not in df["email"].values:
        df.loc[len(df)] = [email, interest]
        df.to_csv(f, index=False)

def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)

def save_csv(path: Path, df: pd.DataFrame):
    df.to_csv(path, index=False)
```

---

## 🔹 `data/company_context.md`

```md
# TribalDesk AI — Context for Assistant
- Mission: Empower Tribal sovereignty through AI tools, funding access, and culturally competent consulting.
- Offerings: Grant proposal generator, policy templates, funding database/tracker, peacekeeping mediation consulting & training.
- Target: Tribal governments/courts, Tribal nonprofits, Native-owned small businesses, grant writers.
- USPs: Tribal focus, cultural competency, time-saving AI, accessible pricing.
- Tone: Respectful, clear, practical. Avoid jargon. Emphasize sovereignty and community impact.
```

---

## 🔹 `requirements.txt`

```
streamlit>=1.36.0
openai>=1.42.0
pandas>=2.0.0
python-docx>=1.0.0
pillow>=10.0.0
```

> (Optional later: `gspread`, `google-auth` for Google Sheets storage.)

---

## 🔹 `.gitignore`

```
# Python
__pycache__/
*.pyc

# Streamlit
.streamlit/

# Data (keep private on GitHub)
data/*.csv
```

---

## 🔹 `README.md`

```md
# TribalDesk AI (Streamlit + OpenAI)

A simple, monetizable app for Tribal governments, nonprofits, and Native-owned businesses.

## Features
- Landing page with email capture + Payhip links
- Chat assistant (OpenAI)
- Proposal generator with .docx export
- Grant finder & tracker (CSV)

## 1) Setup
1. **Create repo**: `tribaldesk-ai` (public or private)
2. Add all files from this pack.
3. Commit and push to GitHub.

## 2) Streamlit Cloud
1. Create a new app from this GitHub repo.
2. Set the **Main file** to `streamlit_app.py`.
3. Add **Secrets** (Settings → Secrets):
```

OPENAI\_API\_KEY = "sk-..."
OPENAI\_MODEL = "gpt-4o-mini"
CONTACT\_EMAIL = "[info@tribaldeskai.com](mailto:info@tribaldeskai.com)"
PAYHIP\_STORE\_URL = "[https://payhip.com/YOUR\_STORE](https://payhip.com/YOUR_STORE)"

```
4. Deploy.

## 3) Customize
- Edit `data/company_context.md` to steer the assistant.
- Change headings/section copy in `streamlit_app.py`.
- Replace Payhip URL in Secrets.

## 4) Monetize
- Add products on Payhip. Put the store or product links in `PAYHIP_STORE_URL`.
- Add more CTA buttons to pages.

## 5) Next Steps
- Add Google Analytics or Plausible via `st.components.v1.iframe`.
- Move email storage to Google Sheets.
```

---

## 🔹 `.streamlit/secrets.toml` (Template — add in Streamlit Cloud UI, *not* in repo)

```toml
OPENAI_API_KEY = "sk-..."
OPENAI_MODEL = "gpt-4o-mini"
CONTACT_EMAIL = "info@tribaldeskai.com"
PAYHIP_STORE_URL = "https://payhip.com/YOUR_STORE"
```

---

## ✅ Notes

* **Never hardcode API keys.** Use Streamlit Secrets.
* CSV files (emails, grants) are stored in `data/` and ignored by Git.
* You can import this repo fresh, delete old ones, and be live in minutes.
* When ready, add a `Pricing` page that links to specific Payhip products.
